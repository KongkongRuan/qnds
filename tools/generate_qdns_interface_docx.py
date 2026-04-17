from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "QDMS集中管理V0.000接口设计.docx"
SOURCE_PATH = ROOT / "QDMS调用QDNS接口说明.md"
STATUS_SOURCE_PATH = ROOT / "QDNS采集状态字段说明.md"
OUTPUT_PATH = ROOT / "QDMS调用QDNS接口设计.docx"
BASE_URL = "http://ip:18023"


def clean_inline_markdown(text: str) -> str:
    cleaned = text.replace("`", "").replace("**", "")
    cleaned = cleaned.replace("&lt;", "<").replace("&gt;", ">")
    return cleaned.strip()


def parse_sections(text: str) -> dict[str, dict[str, object]]:
    sections: dict[str, dict[str, object]] = {}
    current: dict[str, object] | None = None
    for raw_line in text.splitlines():
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", raw_line)
        if heading_match:
            current = {
                "level": len(heading_match.group(1)),
                "title": heading_match.group(2).strip(),
                "lines": [],
            }
            sections[current["title"]] = current
            continue
        if current is not None:
            current["lines"].append(raw_line)
    return sections


def is_table_divider(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return bool(cells) and all(cell and set(cell) <= {"-", ":", " "} for cell in cells)


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return lines[index].strip().startswith("|") and is_table_divider(lines[index + 1])


def parse_markdown_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [clean_inline_markdown(cell) for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        raise ValueError("Markdown table is incomplete.")
    return rows[0], rows[2:]


def parse_blocks(lines: list[str]) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            if index < len(lines):
                index += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines).rstrip()})
            continue

        if is_table_start(lines, index):
            table_lines = [lines[index], lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            headers, rows = parse_markdown_table(table_lines)
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        text_lines = []
        while index < len(lines):
            current = lines[index]
            current_stripped = current.strip()
            if not current_stripped or current_stripped.startswith("```") or is_table_start(lines, index):
                break
            text_lines.append(current.rstrip())
            index += 1
        blocks.append({"type": "text", "text": "\n".join(text_lines).strip()})

    return blocks


def normalize_for_match(text: str) -> str:
    return re.sub(r"\s+", "", clean_inline_markdown(text))


def find_text_metadata(section_lines: list[str]) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for line in section_lines:
        match = re.match(r"^\s*-\s*([^：]+)：(.*)$", line.strip())
        if not match:
            continue
        key = clean_inline_markdown(match.group(1))
        value = clean_inline_markdown(match.group(2))
        metadata[key] = value
    return metadata


def find_table_after_label(blocks: list[dict[str, object]], label_fragment: str) -> tuple[list[str], list[list[str]]] | None:
    target = normalize_for_match(label_fragment)
    for index in range(len(blocks) - 1):
        current = blocks[index]
        next_block = blocks[index + 1]
        if current["type"] == "text" and next_block["type"] == "table":
            if target in normalize_for_match(str(current["text"])):
                return list(next_block["headers"]), list(next_block["rows"])
    return None


def find_code_after_label(blocks: list[dict[str, object]], label_fragment: str) -> tuple[str, str] | None:
    target = normalize_for_match(label_fragment)
    for index in range(len(blocks) - 1):
        current = blocks[index]
        next_block = blocks[index + 1]
        if current["type"] == "text" and next_block["type"] == "code":
            if target in normalize_for_match(str(current["text"])):
                label = clean_inline_markdown(str(current["text"]).rstrip("：:"))
                return label, str(next_block["text"])
    return None


def find_first_table(section_lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    for block in parse_blocks(section_lines):
        if block["type"] == "table":
            return list(block["headers"]), list(block["rows"])
    return None


def extract_bullet_items(section_lines: list[str]) -> list[str]:
    items = []
    for line in section_lines:
        match = re.match(r"^\s*-\s+(.*)$", line)
        if match:
            items.append(clean_inline_markdown(match.group(1)))
    return items


def rename_headers(headers: list[str], target_headers: list[str]) -> list[str]:
    if len(headers) != len(target_headers):
        return headers
    return target_headers


def get_section(sections: dict[str, dict[str, object]], title: str) -> list[str]:
    return list(sections[title]["lines"])


def build_sections() -> tuple[dict[str, dict[str, object]], dict[str, dict[str, object]]]:
    source_sections = parse_sections(SOURCE_PATH.read_text(encoding="utf-8"))
    status_sections = parse_sections(STATUS_SOURCE_PATH.read_text(encoding="utf-8"))
    return source_sections, status_sections


def build_generic_response_rows(data_desc: str) -> list[list[str]]:
    return [
        ["code", "int", "业务状态码，200 表示成功"],
        ["message", "String", "执行结果描述"],
        ["data", "Object / Array", data_desc],
    ]


def clear_document_body(document: Document) -> None:
    body = document._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_run_font(run, font_name: str, size_pt: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_section_label(document: Document, text: str) -> None:
    add_paragraph(document, text)


def add_paragraph(document: Document, text: str = "", style: str = "Normal") -> None:
    paragraph = document.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)


def add_label_value(document: Document, label: str, value: str) -> None:
    add_paragraph(document, f"{label}: {value}")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.style = "Normal"
        for run in paragraph.runs:
            run.bold = bold


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(row_cells[index], value)


def add_code_block(document: Document, code_text: str) -> None:
    lines = code_text.splitlines() or [""]
    for line in lines:
        paragraph = document.add_paragraph(style="HTML Preformatted")
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, "Consolas", size_pt=9)


def add_numbered_notes(document: Document, notes: Iterable[str]) -> None:
    cleaned_notes = [clean_inline_markdown(note) for note in notes if clean_inline_markdown(note)]
    if not cleaned_notes:
        return
    add_paragraph(document, "说明:")
    for index, note in enumerate(cleaned_notes, start=1):
        add_paragraph(document, f"{index}. {note}")


def render_examples(document: Document, examples: list[tuple[str, str]]) -> None:
    for label, code in examples:
        add_paragraph(document, f"{label}:")
        add_code_block(document, code)


def filter_examples(examples: Iterable[tuple[str, str] | None]) -> list[tuple[str, str]]:
    return [example for example in examples if example]


def render_interface(
    document: Document,
    *,
    title: str,
    method: str,
    path: str,
    request_headers: list[str] | None = None,
    request_rows: list[list[str]] | None = None,
    request_text: str | None = None,
    response_headers: list[str] | None = None,
    response_rows: list[list[str]] | None = None,
    response_text: str | None = None,
    request_examples: list[tuple[str, str]] | None = None,
    response_examples: list[tuple[str, str]] | None = None,
    notes: list[str] | None = None,
) -> None:
    add_heading(document, title, 2)
    add_label_value(document, "请求URL", f"{BASE_URL}{path}")
    add_label_value(document, "Method", method)
    if method.upper() == "POST":
        add_label_value(document, "Content-Type", "application/json")

    if request_headers and request_rows is not None:
        add_paragraph(document, "请求参数:")
        add_table(document, request_headers, request_rows)
    else:
        add_paragraph(document, f"请求参数: {request_text or '无'}")

    if response_headers and response_rows is not None:
        add_paragraph(document, "响应:")
        add_table(document, response_headers, response_rows)
    else:
        add_paragraph(document, f"响应: {response_text or '参见下方响应示例。'}")

    if request_examples:
        render_examples(document, request_examples)
    if response_examples:
        render_examples(document, response_examples)
    if notes:
        add_numbered_notes(document, notes)
    add_paragraph(document)


def render_intro_module(document: Document) -> None:
    add_heading(document, "QDMS调用QDNS接口设计文档", 1)
    add_heading(document, "1. 说明", 1)
    add_paragraph(document, "本文档用于说明 QDMS 平台调用 QDNS 节点时涉及的北向接口。当前调用链路为：QDMS -> QDNS -> VPN-Sim。")
    add_table(
        document,
        ["项目", "说明"],
        [
            ["默认端口", "18023"],
            ["GET 调用约定", "固定路径 + 查询参数，不使用 REST 风格路径变量传参"],
            ["POST 调用约定", "固定路径 + JSON 请求体传参，不把业务参数放到 URI 路径中"],
            ["设备标识约定", "QDMS 传给 QDNS 的 deviceId 表示 QDMS 平台设备 ID，不要求与 VPN-Sim 中的 device_id 一致"],
            ["节点寻址方式", "QDNS 收到运维请求后，会先根据 deviceId 在节点缓存 / Redis 设备列表中查找 deviceIp、devicePort，再调用下游设备"],
            ["统一运维入口", "/api/unified/*"],
            ["证书编排入口", "/api/device-certificates/*"],
        ],
    )
    add_paragraph(document)
    add_heading(document, "2. 基础信息", 1)
    add_section_label(document, "返回约定:")
    add_table(
        document,
        ["返回风格", "示例", "说明"],
        [
            ["ApiResponse", '{ "code": 200, "message": "success", "data": ... }', "统一运维和证书编排接口主要使用该包装结构"],
            ["普通 Map", '{ "code": 200, "message": "...", "data": ... }', "部分节点接口使用普通 Map 返回，建议同样按 code == 200 判断成功"],
        ],
    )
    add_paragraph(document, "建议 QDMS 统一按 code == 200 判断 HTTP 业务成功，再结合 data.status 或 task.status 判断运维任务结果。")
    add_paragraph(document)


def render_node_info_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    section_lines = get_section(source_sections, "3.1 获取节点信息")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    response_example = find_code_after_label(blocks, "返回示例")
    response_rows = [
        ["hostname", "String", "节点主机名"],
        ["ip", "String", "节点 IP 地址"],
        ["appName", "String", "应用名称"],
        ["port", "int", "服务监听端口"],
        ["nodeKey", "String", "节点唯一标识，通常为 ip:port"],
        ["javaVersion", "String", "JVM 版本"],
        ["osName", "String", "操作系统名称"],
        ["uptime", "String", "应用运行时长"],
    ]
    add_heading(document, "3. 节点信息接口", 1)
    render_interface(
        document,
        title="3.1 获取节点信息",
        method=metadata["方法"],
        path=metadata["路径"],
        request_text="无",
        response_headers=["字段名", "类型", "说明"],
        response_rows=response_rows,
        response_examples=filter_examples([response_example]),
        notes=[metadata.get("说明", "用于获取节点 IP、端口、节点标识以及 JVM 运行信息。")],
    )


def render_device_sync_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    section_lines = get_section(source_sections, "4.1 向节点同步单台设备")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    request_table = find_table_after_label(blocks, "请求体字段")
    request_example = find_code_after_label(blocks, "请求示例")
    response_example = find_code_after_label(blocks, "返回示例")
    response_rows = [
        ["code", "int", "业务状态码，200 表示设备同步成功"],
        ["message", "String", "执行结果描述"],
        ["trapTargetSet", "boolean", "是否同步设置 Trap 目标"],
        ["data.id", "String", "节点侧设备标识"],
        ["data.name", "String", "设备名称"],
        ["data.deviceIp", "String", "设备 IP"],
        ["data.devicePort", "String", "设备 SNMP 端口"],
        ["data.protocol", "String", "SNMP 协议版本"],
    ]
    add_heading(document, "4. 设备同步接口", 1)
    render_interface(
        document,
        title="4.1 向节点同步单台设备",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=rename_headers(list(request_table[0]), ["参数名", "是否必填", "说明"]) if request_table else None,
        request_rows=list(request_table[1]) if request_table else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=response_rows,
        request_examples=filter_examples([request_example]),
        response_examples=filter_examples([response_example]),
        notes=[metadata.get("说明", "QDMS 将设备下发给指定 QDNS 节点，由节点建立本地缓存并纳入采集队列。")],
    )


def render_device_query_collect_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    add_heading(document, "5. 设备查询与采集接口", 1)

    section_lines = get_section(source_sections, "5.1 获取节点设备列表")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    request_table = find_table_after_label(blocks, "查询参数")
    render_interface(
        document,
        title="5.1 获取节点设备列表",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=rename_headers(list(request_table[0]), ["参数名", "是否必填", "说明"]) if request_table else None,
        request_rows=list(request_table[1]) if request_table else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("设备列表或分页结果；当 page/size 不传或 size <= 0 时返回全部设备。"),
        notes=[metadata.get("说明", "用于查询节点当前持有的设备列表。")],
    )

    section_lines = get_section(source_sections, "5.2 获取单台设备最新状态")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    request_table = find_table_after_label(blocks, "查询参数")
    render_interface(
        document,
        title="5.2 获取单台设备最新状态",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=rename_headers(list(request_table[0]), ["参数名", "是否必填", "说明"]) if request_table else None,
        request_rows=list(request_table[1]) if request_table else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("单台设备最近一次采集到的状态对象，字段见后文“设备状态字段说明”。"),
        notes=[metadata.get("说明", "返回节点最近一次采集到的状态。")],
    )

    section_lines = get_section(source_sections, "5.3 获取全部设备最新状态")
    metadata = find_text_metadata(section_lines)
    render_interface(
        document,
        title="5.3 获取全部设备最新状态",
        method=metadata["方法"],
        path=metadata["路径"],
        request_text="无",
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("节点缓存中的全部设备状态列表，单个元素字段见后文“设备状态字段说明”。"),
        notes=[metadata.get("说明", "返回节点缓存中的所有设备状态。")],
    )

    section_lines = get_section(source_sections, "5.4 触发单台设备采集")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    request_example = find_code_after_label(blocks, "请求示例")
    render_interface(
        document,
        title="5.4 触发单台设备采集",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=["参数名", "是否必填", "说明"],
        request_rows=[
            ["id", "否", "节点侧设备标识；与 deviceId 二选一传入"],
            ["deviceId", "否", "QDMS 平台设备 ID；与 id 二选一传入"],
        ],
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("本次采集任务的受理结果或采集结果摘要。"),
        request_examples=filter_examples([request_example]),
        notes=[
            metadata.get("说明", "手动触发单台设备采集。"),
            "常用于运维完成后刷新版本号、在线状态和证书状态。",
        ],
    )

    section_lines = get_section(source_sections, "5.5 触发全量采集")
    metadata = find_text_metadata(section_lines)
    render_interface(
        document,
        title="5.5 触发全量采集",
        method=metadata["方法"],
        path=metadata["路径"],
        request_text="无",
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("全量采集任务的受理结果或执行摘要。"),
        notes=[metadata.get("说明", "手动触发节点对全部设备做一次采集。")],
    )


def render_status_fields_module(document: Document, status_sections: dict[str, dict[str, object]]) -> None:
    add_heading(document, "5.6 设备状态字段说明", 2)
    add_paragraph(document, "本章节适用于 /api/device/status 与 /api/device/status/all 返回中的 data 对象说明。")
    add_paragraph(document)

    top_level_table = find_first_table(get_section(status_sections, "3. 顶层字段说明"))
    if top_level_table:
        add_section_label(document, "顶层状态字段:")
        add_table(document, list(top_level_table[0]), list(top_level_table[1]))
        add_paragraph(document)

    interface_table = find_first_table(get_section(status_sections, "4.1 单个接口对象字段"))
    if interface_table:
        add_section_label(document, "接口对象字段:")
        add_table(document, list(interface_table[0]), list(interface_table[1]))
        add_paragraph(document)

    tunnel_table = find_first_table(get_section(status_sections, "5.1 单条隧道对象字段"))
    if tunnel_table:
        add_section_label(document, "隧道对象字段:")
        add_table(document, list(tunnel_table[0]), list(tunnel_table[1]))
        add_paragraph(document)

    ike_table = find_first_table(get_section(status_sections, "6.1 单条 IKE SA 对象字段"))
    if ike_table:
        add_section_label(document, "IKE SA 对象字段:")
        add_table(document, list(ike_table[0]), list(ike_table[1]))
        add_paragraph(document)

    ipsec_table = find_first_table(get_section(status_sections, "7.1 单条 IPsec SA 对象字段"))
    if ipsec_table:
        add_section_label(document, "IPsec SA 对象字段:")
        add_table(document, list(ipsec_table[0]), list(ipsec_table[1]))
        add_paragraph(document)


def render_log_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    add_heading(document, "6. 日志拉取接口", 1)

    section_lines = get_section(source_sections, "6.1 拉取单台设备日志")
    metadata = find_text_metadata(section_lines)
    request_example = find_code_after_label(parse_blocks(section_lines), "请求示例")
    render_interface(
        document,
        title="6.1 拉取单台设备日志",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=["参数名", "是否必填", "说明"],
        request_rows=[
            ["deviceId", "是", "目标设备的 QDMS 设备 ID"],
            ["count", "否", "拉取的日志条数，默认由节点侧决定"],
            ["since", "否", "增量拉取起始时间；传 null 表示不按时间过滤"],
        ],
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("单台设备日志列表。"),
        request_examples=filter_examples([request_example]),
        notes=["用于拉取指定设备最近日志，可按条数或时间做增量过滤。"],
    )

    section_lines = get_section(source_sections, "6.2 批量拉取指定设备日志")
    metadata = find_text_metadata(section_lines)
    request_example = find_code_after_label(parse_blocks(section_lines), "请求示例")
    render_interface(
        document,
        title="6.2 批量拉取指定设备日志",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=["参数名", "是否必填", "说明"],
        request_rows=[
            ["deviceIds", "是", "目标设备 ID 列表"],
            ["count", "否", "每台设备拉取的日志条数"],
        ],
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("按设备聚合后的日志结果。"),
        request_examples=filter_examples([request_example]),
        notes=["用于批量拉取指定设备的最近日志。"],
    )

    section_lines = get_section(source_sections, "6.3 拉取全部设备日志")
    metadata = find_text_metadata(section_lines)
    request_example = find_code_after_label(parse_blocks(section_lines), "请求示例")
    render_interface(
        document,
        title="6.3 拉取全部设备日志",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=["参数名", "是否必填", "说明"],
        request_rows=[["count", "否", "每台设备需要拉取的日志条数"]],
        response_headers=["字段名", "类型", "说明"],
        response_rows=build_generic_response_rows("节点内全部设备的日志结果。"),
        request_examples=filter_examples([request_example]),
        notes=["用于一次性拉取节点当前全部设备日志。"],
    )


def render_unified_command_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    add_heading(document, "7. 统一运维接口", 1)

    section_lines = get_section(source_sections, "7.1 提交统一运维命令")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    common_request = find_code_after_label(blocks, "通用请求体")
    common_response = find_code_after_label(blocks, "通用返回体")
    backup_request = find_code_after_label(parse_blocks(get_section(source_sections, "9.1 导出备份")), "请求")
    backup_downstream = find_code_after_label(parse_blocks(get_section(source_sections, "9.1 导出备份")), "downstream")
    restore_request = find_code_after_label(parse_blocks(get_section(source_sections, "9.2 恢复备份")), "请求")
    upgrade_request = find_code_after_label(parse_blocks(get_section(source_sections, "9.3 升级固件")), "请求")
    reboot_request = find_code_after_label(parse_blocks(get_section(source_sections, "9.4 重启设备")), "请求")

    render_interface(
        document,
        title="7.1 提交统一运维命令",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=["参数名", "是否必填", "说明"],
        request_rows=[
            ["requestId", "否", "请求流水号；不传时可由 QDNS 自动生成"],
            ["deviceId", "是", "QDMS 平台设备 ID"],
            ["code", "是", "统一运维能力编码"],
            ["operation", "是", "具体操作名，应与 code 对应"],
            ["payload", "否", "业务参数对象；不同操作的必填字段见“统一运维能力编码说明”"],
            ["operator", "否", "操作人或调用方标识，例如 qdms"],
        ],
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码，200 表示请求已受理"],
            ["message", "String", "执行结果描述"],
            ["data.requestId", "String", "请求流水号"],
            ["data.taskId", "String", "统一运维任务 ID"],
            ["data.deviceId", "String", "目标设备 ID"],
            ["data.code", "int", "统一运维能力编码"],
            ["data.operation", "String", "具体操作名"],
            ["data.mode", "String", "执行模式，SYNC 或 ASYNC"],
            ["data.status", "String", "当前任务状态，如 ACCEPTED / SUCCESS / FAILED"],
            ["data.statusMessage", "String", "状态描述"],
            ["data.downstream", "Object / null", "下游设备返回结果；异步受理时通常为 null"],
        ],
        request_examples=filter_examples(
            [
                common_request,
                ("导出备份请求示例", backup_request[1]) if backup_request else None,
                ("恢复备份请求示例", restore_request[1]) if restore_request else None,
                ("升级固件请求示例", upgrade_request[1]) if upgrade_request else None,
                ("重启设备请求示例", reboot_request[1]) if reboot_request else None,
            ]
        ),
        response_examples=filter_examples(
            [
                common_response,
                ("导出备份成功响应中的 downstream 示例", backup_downstream[1]) if backup_downstream else None,
            ]
        ),
        notes=[
            "10005/10006/10007/10008 已改为由 QDNS 通过 HTTP 调用 VPN-Sim 的设备接口。",
            "upgrade 与 reboot 为异步任务；调用方需要结合 mode 和 status 判断任务是否真正完成。",
        ],
    )

    section_lines = get_section(source_sections, "7.2 查询统一运维异步任务")
    metadata = find_text_metadata(section_lines)
    blocks = parse_blocks(section_lines)
    request_table = find_table_after_label(blocks, "查询参数")
    response_example = find_code_after_label(blocks, "返回示例")
    render_interface(
        document,
        title="7.2 查询统一运维异步任务",
        method=metadata["方法"],
        path=metadata["路径"],
        request_headers=rename_headers(list(request_table[0]), ["参数名", "是否必填", "说明"]) if request_table else None,
        request_rows=list(request_table[1]) if request_table else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data.taskId", "String", "统一运维任务 ID"],
            ["data.requestId", "String", "请求流水号"],
            ["data.deviceId", "String", "目标设备 ID"],
            ["data.code", "int", "统一运维能力编码"],
            ["data.operation", "String", "具体操作名"],
            ["data.mode", "String", "执行模式"],
            ["data.status", "String", "任务状态，如 ACCEPTED / RUNNING / SUCCESS / FAILED"],
            ["data.statusMessage", "String", "状态描述"],
            ["data.downstream", "Object / null", "最终下游结果；异步任务成功后通常会补充"],
            ["data.createdAt", "String", "任务创建时间"],
            ["data.startedAt", "String", "任务开始时间"],
            ["data.finishedAt", "String", "任务结束时间"],
        ],
        response_examples=filter_examples([response_example]),
        notes=["对于 ASYNC 操作，平台必须继续轮询该接口，不要只看受理返回。"],
    )


def render_capability_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    add_heading(document, "8. 统一运维能力编码表", 1)
    add_paragraph(document, "以下能力均通过 /api/unified/command 提交，payload 字段按下表中的必填字段组织。")
    add_paragraph(document)

    ssl_table = find_first_table(get_section(source_sections, "8.1 SSL VPN 配置类"))
    if ssl_table:
        add_heading(document, "8.1 SSL VPN 配置类", 2)
        add_table(document, list(ssl_table[0]), list(ssl_table[1]))
        add_paragraph(document)

    ops_table = find_first_table(get_section(source_sections, "8.2 设备运维类"))
    if ops_table:
        add_heading(document, "8.2 设备运维类", 2)
        add_table(document, list(ops_table[0]), list(ops_table[1]))
        add_paragraph(document)

    cert_table = find_first_table(get_section(source_sections, "8.3 证书类"))
    if cert_table:
        add_heading(document, "8.3 证书原子能力", 2)
        add_table(document, list(cert_table[0]), list(cert_table[1]))
        cert_notes = extract_bullet_items(get_section(source_sections, "8.3 证书类"))
        add_numbered_notes(document, cert_notes)
        add_paragraph(document)


def build_certificate_table_from_8_4(
    source_sections: dict[str, dict[str, object]],
    label_fragment: str,
) -> tuple[list[str], list[list[str]]] | None:
    section_lines = get_section(source_sections, "8.4 证书编排接口")
    table = find_table_after_label(parse_blocks(section_lines), label_fragment)
    if not table:
        return None
    return rename_headers(list(table[0]), ["参数名", "是否必填", "说明"]), list(table[1])


def render_certificate_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    add_section_label(document, "证书编排接口说明:")
    add_paragraph(document, "本模块路径前缀为 /api/device-certificates，统一返回 ApiResponse。")
    add_paragraph(document, "适用场景包括证书状态同步、证书读取、CSR 生成 / 回读、自建 CA 签发并自动下发，以及第三方证书安装。")
    add_paragraph(document)

    status_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/status")
    render_interface(
        document,
        title="8.4 查询设备证书状态",
        method="GET",
        path="/api/device-certificates/status",
        request_headers=status_request[0] if status_request else None,
        request_rows=status_request[1] if status_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data", "Object", "设备证书状态聚合结果，包含设备证书与 CA 链安装状态摘要"],
        ],
        notes=["用于查询设备当前证书状态。"],
    )

    sync_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/sync")
    render_interface(
        document,
        title="8.5 同步设备当前证书信息",
        method="POST",
        path="/api/device-certificates/sync",
        request_headers=sync_request[0] if sync_request else None,
        request_rows=sync_request[1] if sync_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data.deviceId", "String", "目标设备 ID"],
            ["data.status", "Object", "聚合后的设备证书状态"],
            ["data.deviceCertificate", "Object", "当前设备证书"],
            ["data.caChain", "Array", "当前 CA 证书链"],
            ["data.csr", "Object", "当前 CSR 信息"],
        ],
        notes=["用于聚合状态、设备证书、CA 链与 CSR 信息，便于前端一次刷新。"],
    )

    device_cert_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/device-cert")
    render_interface(
        document,
        title="8.6 读取设备证书",
        method="GET",
        path="/api/device-certificates/device-cert",
        request_headers=device_cert_request[0] if device_cert_request else None,
        request_rows=device_cert_request[1] if device_cert_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data", "Object", "当前设备已安装证书信息"],
        ],
        notes=["用于读取设备当前已安装的设备证书。"],
    )

    ca_chain_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/ca-chain")
    render_interface(
        document,
        title="8.7 读取 CA 证书链",
        method="GET",
        path="/api/device-certificates/ca-chain",
        request_headers=ca_chain_request[0] if ca_chain_request else None,
        request_rows=ca_chain_request[1] if ca_chain_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data", "Array / Object", "当前设备已安装的 CA 证书链信息"],
        ],
        notes=["用于读取设备当前已安装的 CA 证书链。"],
    )

    csr_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/csr")
    csr_section = get_section(source_sections, "9.5 查询或生成设备 CSR")
    csr_blocks = parse_blocks(csr_section)
    csr_request_example = find_code_after_label(csr_blocks, "查询参数示例")
    csr_response_example = find_code_after_label(csr_blocks, "返回示例")
    render_interface(
        document,
        title="8.8 读取或生成设备 CSR",
        method="GET",
        path="/api/device-certificates/csr",
        request_headers=csr_request[0] if csr_request else None,
        request_rows=csr_request[1] if csr_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data.deviceId", "String", "目标设备 ID"],
            ["data.csr.code", "int", "证书原子能力编码，固定为 10014"],
            ["data.csr.operation", "String", "原子能力操作名，固定为 read_device_csr"],
            ["data.csr.mode", "String", "执行模式，通常为 SYNC"],
            ["data.csr.status", "String", "步骤状态"],
            ["data.csr.statusMessage", "String", "步骤状态描述"],
            ["data.csr.downstream.algorithm", "String", "CSR 使用的算法"],
            ["data.csr.downstream.csrPemBase64", "String", "CSR PEM 文本的 Base64 编码"],
        ],
        request_examples=filter_examples([csr_request_example]),
        response_examples=filter_examples([csr_response_example]),
        notes=["forceRegenerate=true 时会强制重新生成 CSR。"],
    )

    self_ca_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/issue/self-ca")
    self_ca_section = get_section(source_sections, "9.6 自建 CA 签发并下发")
    self_ca_blocks = parse_blocks(self_ca_section)
    self_ca_request_example = find_code_after_label(self_ca_blocks, "请求")
    self_ca_response_example = find_code_after_label(self_ca_blocks, "返回示例")
    render_interface(
        document,
        title="8.9 自建 CA 签发并下发",
        method="POST",
        path="/api/device-certificates/issue/self-ca",
        request_headers=self_ca_request[0] if self_ca_request else None,
        request_rows=self_ca_request[1] if self_ca_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data.device", "Object", "目标设备基础信息，仅保留展示和定位所需字段"],
            ["data.deviceCertificate", "Object", "安装完成后最终返回并回读确认的设备证书信息"],
            ["data.caSerialNumber", "String", "CA 平台中的证书序列号，采用十进制字符串表示"],
        ],
        request_examples=filter_examples([self_ca_request_example]),
        response_examples=filter_examples([self_ca_response_example]),
        notes=[
            "接口内部会执行读取 CSR、CA 签发、获取 CA 链、下发以及回读校验整套流程。",
            "旧的详细编排结果仍保留在服务内部逻辑中，但不再作为 HTTP 成功响应直接返回。",
            "caSerialNumber 采用 CA 平台使用口径的十进制序列号；如果设备侧展示为十六进制，两边表现形式可能不同，但对应同一张证书。",
        ],
    )

    third_party_request = build_certificate_table_from_8_4(source_sections, "/api/device-certificates/install/third-party")
    third_party_section = get_section(source_sections, "9.7 安装第三方证书")
    third_party_blocks = parse_blocks(third_party_section)
    third_party_request_example = find_code_after_label(third_party_blocks, "请求")
    render_interface(
        document,
        title="8.10 安装第三方证书",
        method="POST",
        path="/api/device-certificates/install/third-party",
        request_headers=third_party_request[0] if third_party_request else None,
        request_rows=third_party_request[1] if third_party_request else None,
        response_headers=["字段名", "类型", "说明"],
        response_rows=[
            ["code", "int", "业务状态码"],
            ["message", "String", "执行结果描述"],
            ["data", "Object", "第三方证书安装结果以及安装后的回读校验信息"],
        ],
        request_examples=filter_examples([third_party_request_example]),
        notes=[
            "该接口会先读取或生成设备 CSR，再下发第三方 CA 链和设备证书。",
            "接口返回中会给出安装结果以及回读校验信息，便于 QDMS 直接展示安装是否成功。",
        ],
    )


def render_examples_summary_module(document: Document) -> None:
    add_heading(document, "9. 运维与证书接口示例", 1)
    add_paragraph(document, "本版文档已将原示例章节中的请求示例、响应示例和说明，分别并入 7.1、8.8、8.9、8.10 等对应接口条目中，便于按接口直接查阅。")
    add_paragraph(document)


def render_suggestion_module(document: Document, source_sections: dict[str, dict[str, object]]) -> None:
    add_heading(document, "10. 对接建议", 1)
    add_section_label(document, "10.1 对接建议:")
    suggestion_rows = []
    for index, item in enumerate(extract_bullet_items(get_section(source_sections, "10. 对接建议")), start=1):
        suggestion_rows.append([str(index), item])
    add_table(document, ["序号", "建议"], suggestion_rows)
    add_paragraph(document)


def render_all(
    document: Document,
    source_sections: dict[str, dict[str, object]],
    status_sections: dict[str, dict[str, object]],
) -> None:
    render_intro_module(document)
    render_node_info_module(document, source_sections)
    render_device_sync_module(document, source_sections)
    render_device_query_collect_module(document, source_sections)
    render_status_fields_module(document, status_sections)
    render_log_module(document, source_sections)
    render_unified_command_module(document, source_sections)
    render_capability_module(document, source_sections)
    render_certificate_module(document, source_sections)
    render_examples_summary_module(document)
    render_suggestion_module(document, source_sections)


def main() -> None:
    source_sections, status_sections = build_sections()
    document = Document(str(TEMPLATE_PATH))
    clear_document_body(document)
    render_all(document, source_sections, status_sections)
    document.save(str(OUTPUT_PATH))
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
